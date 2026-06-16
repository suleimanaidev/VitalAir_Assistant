"""Extract readable text from patient health uploads (PDF, Word, images, text)."""

from __future__ import annotations

import re
from collections.abc import Callable
from io import BytesIO

from services.ocr_setup import configure_tesseract, ocr_unavailable_message

ALLOWED_EXTENSIONS = {
    ".txt",
    ".md",
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".webp",
    ".doc",
    ".docx",
}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}
IMAGE_CONTENT_TYPES = {"image/png", "image/jpeg", "image/jpg", "image/webp"}
WORD_EXTENSIONS = {".doc", ".docx"}
WORD_CONTENT_TYPES = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def _normalize_whitespace(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _decode_text_file(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return _normalize_whitespace(raw.decode(encoding))
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not read text file encoding")


def _extract_docx(raw: bytes) -> str:
    from docx import Document

    doc = Document(BytesIO(raw))
    parts: list[str] = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            parts.append(line)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                line = cell.text.strip()
                if line:
                    parts.append(line)
    return _normalize_whitespace("\n".join(parts))


def _extract_doc_legacy(raw: bytes) -> str:
    """Best-effort text from old binary .doc (Word 97–2003) via OLE stream."""
    try:
        import olefile
    except ImportError:
        return ""

    try:
        ole = olefile.OleFileIO(BytesIO(raw))
    except Exception:
        return ""

    if not ole.exists("WordDocument"):
        return ""

    data = ole.openstream("WordDocument").read()
    chars: list[str] = []
    i = 0
    while i < len(data) - 1:
        pair = data[i : i + 2]
        if pair == b"\x00\x00":
            i += 2
            continue
        try:
            ch = pair.decode("utf-16-le")
        except UnicodeDecodeError:
            i += 1
            continue
        if ch.isprintable() or ch in "\n\r\t":
            chars.append(ch)
        i += 2

    text = _normalize_whitespace("".join(chars))
    if len(text) >= 20:
        return text

    ascii_runs = re.findall(rb"[\x20-\x7e\n\r\t]{8,}", data)
    fallback = _normalize_whitespace(
        b" ".join(ascii_runs).decode("ascii", errors="ignore")
    )
    return fallback if len(fallback) >= 20 else ""


def _extract_pdf_pypdf(raw: bytes, max_pages: int = 20) -> str:
    from pypdf import PdfReader  # type: ignore[import-untyped]

    reader = PdfReader(BytesIO(raw))
    pages = [p.extract_text() or "" for p in reader.pages[:max_pages]]
    return _normalize_whitespace("\n".join(pages))


def _extract_pdf_pdfplumber(raw: bytes, max_pages: int = 20) -> str:
    import pdfplumber

    texts: list[str] = []
    with pdfplumber.open(BytesIO(raw)) as pdf:
        for page in pdf.pages[:max_pages]:
            texts.append(page.extract_text() or "")
    return _normalize_whitespace("\n".join(texts))


def _extract_pdf_pymupdf(raw: bytes, max_pages: int = 20) -> str:
    import fitz

    doc = fitz.open(stream=raw, filetype="pdf")
    texts = [str(page.get_text() or "") for page in doc[:max_pages]]
    return _normalize_whitespace("\n".join(texts))


def _prepare_image_for_ocr(img):
    from PIL import Image, ImageEnhance, ImageOps

    img = ImageOps.exif_transpose(img)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    gray = img.convert("L")
    w, h = gray.size
    if max(w, h) < 1400:
        scale = 1400 / max(w, h)
        try:
            resample = Image.Resampling.LANCZOS  # type: ignore[attr-defined]
        except AttributeError:
            resample = Image.LANCZOS  # type: ignore[attr-defined]
        gray = gray.resize((int(w * scale), int(h * scale)), resample=resample)
    return ImageEnhance.Contrast(gray).enhance(1.4)


def _ocr_tesseract(img) -> str:
    import pytesseract

    if not configure_tesseract():
        return ""
    configs = ["", "--psm 6", "--psm 4", "--psm 11"]
    best = ""
    for cfg in configs:
        try:
            chunk = pytesseract.image_to_string(img, config=cfg).strip()
            if len(chunk) > len(best):
                best = chunk
        except Exception:
            continue
    return _normalize_whitespace(best)


def _ocr_rapid(raw: bytes) -> str:
    try:
        from rapidocr_onnxruntime import RapidOCR  # type: ignore[import-untyped]
    except ImportError:
        return ""

    try:
        engine = RapidOCR()
        result, _ = engine(raw)
        if not result:
            return ""
        lines = [str(line[1]).strip() for line in result if len(line) > 1 and line[1]]
        return _normalize_whitespace("\n".join(lines))
    except Exception:
        return ""


def _ocr_image_bytes(raw: bytes) -> tuple[str, str]:
    """
    OCR image bytes. Returns (text, method).
    Tries Tesseract first, then RapidOCR (no system install needed).
    """
    try:
        from PIL import Image
    except ImportError:
        return "", ""

    try:
        img = _prepare_image_for_ocr(Image.open(BytesIO(raw)))
    except Exception:
        return "", ""

    tesseract_text = _ocr_tesseract(img)
    if len(tesseract_text) >= 12:
        return tesseract_text, "ocr_image"

    rapid_text = _ocr_rapid(raw)
    if len(rapid_text) >= 12:
        return rapid_text, "ocr_rapid"

    if tesseract_text:
        return tesseract_text, "ocr_image"
    if rapid_text:
        return rapid_text, "ocr_rapid"
    return "", ""


def _ocr_pdf_pages(raw: bytes, max_pages: int = 5) -> tuple[str, str]:
    try:
        import fitz
        from PIL import Image
    except ImportError:
        return "", ""

    texts: list[str] = []
    method = ""
    try:
        doc = fitz.open(stream=raw, filetype="pdf")
        for page in doc[:max_pages]:
            pix = page.get_pixmap(dpi=180)
            png = pix.tobytes("png")
            chunk, page_method = _ocr_image_bytes(png)
            if chunk:
                texts.append(chunk)
                method = page_method or "ocr_pdf"
    except Exception:
        return "", ""
    combined = _normalize_whitespace("\n".join(texts))
    return combined, method or "ocr_pdf"


def extract_document_text(
    filename: str,
    content_type: str,
    raw: bytes,
) -> tuple[str, str]:
    """
    Parse patient document bytes → (text, method).
    Chain: text → Word → PDF text layer → OCR (images / scanned PDF).
    """
    if not raw:
        raise ValueError("File is empty")

    lower = filename.lower()
    ext = "." + lower.rsplit(".", 1)[-1] if "." in lower else ""

    if ext in (".txt", ".md") or content_type.startswith("text/"):
        return _decode_text_file(raw), "text"

    if ext in IMAGE_EXTENSIONS or content_type in IMAGE_CONTENT_TYPES:
        text, ocr_method = _ocr_image_bytes(raw)
        if len(text) >= 12:
            return text, ocr_method or "ocr_image"
        raise ValueError(
            f"Could not read text from image. {ocr_unavailable_message()} "
            "Try a clearer, well-lit photo or upload PDF/Word instead."
        )

    if ext in WORD_EXTENSIONS or content_type in WORD_CONTENT_TYPES:
        if ext == ".docx" or "wordprocessingml" in content_type:
            try:
                text = _extract_docx(raw)
                if len(text) >= 15:
                    return text, "docx"
            except ImportError:
                raise ValueError(
                    "Word .docx support requires python-docx. "
                    "Run: pip install python-docx"
                ) from None
            except Exception as exc:
                raise ValueError(f"Could not read .docx file: {exc}") from exc

        legacy = _extract_doc_legacy(raw)
        if len(legacy) >= 15:
            return legacy, "doc_legacy"

        if ext == ".docx":
            raise ValueError(
                "Could not extract text from .docx. "
                "Try exporting as PDF or plain text."
            )
        raise ValueError(
            "Could not read legacy .doc file. Save as .docx or PDF, "
            "or upload a clear photo for OCR."
        )

    if ext == ".pdf" or content_type == "application/pdf":
        methods: list[tuple[str, Callable[[bytes], str]]] = []

        try:
            from pypdf import PdfReader  # type: ignore[import-untyped]  # noqa: F401

            methods.append(("pypdf", _extract_pdf_pypdf))
        except ImportError:
            pass

        try:
            import pdfplumber  # noqa: F401

            methods.append(("pdfplumber", _extract_pdf_pdfplumber))
        except ImportError:
            pass

        try:
            import fitz  # noqa: F401

            methods.append(("pymupdf", _extract_pdf_pymupdf))
        except ImportError:
            pass

        for name, fn in methods:
            try:
                text = fn(raw)
                if len(text) >= 15:
                    return text, name
            except Exception:
                continue

        ocr_text, ocr_method = _ocr_pdf_pages(raw)
        if len(ocr_text) >= 12:
            return ocr_text, ocr_method or "ocr_pdf"

        raise ValueError(
            "PDF has no readable text layer. For scanned prescriptions, "
            f"{ocr_unavailable_message()}"
        )

    raise ValueError(
        "Allowed formats: .pdf, .doc, .docx, .png, .jpg, .jpeg, .webp, .txt, .md"
    )
